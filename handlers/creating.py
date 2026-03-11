from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

def rubles_to_text(amount):
    return num2words(amount, lang='ru', to='currency', currency='RUB').capitalize()

def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    new_text = full_text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != full_text:
        style = None
        if paragraph.runs:
            style = paragraph.runs[0].font
        paragraph.clear()
        run = paragraph.add_run(new_text)
        if style:
            run.font.name = style.name
            run.font.size = style.size
            run.font.bold = style.bold
            run.font.italic = style.italic

def replace_text_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

def find_table_by_headers(doc, headers):
    for table in doc.tables:
        first_row = table.rows[0]
        row_text = ' '.join(cell.text for cell in first_row.cells)
        if all(header in row_text for header in headers):
            return table
    return None

def set_cell_font(cell, font_name, font_size, alignment=None):
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

def set_document_font(doc, font_name):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name

def set_paragraph_font(paragraph, font_name, font_size, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic

def fill_invoice(template_path, output_path, data):
    doc = Document(template_path)
    req_table = None
    for table in doc.tables:
        if any("ИНН" in cell.text for row in table.rows for cell in row.cells):
            req_table = table
            break

    if req_table:
        for row in req_table.rows:
            if len(row.cells) >= 2:
                left = row.cells[0].text
                right = row.cells[1].text

                if "ORG" in right:
                    replace_text_in_cell(row.cells[1], {"ORG": data['buyer_org']})
                elif "INN" in right:
                    replace_text_in_cell(row.cells[1], {"INN": data['buyer_inn']})
                elif "ORGN" in right:
                    replace_text_in_cell(row.cells[1], {"ORGN": data['buyer_ogrn']})
                elif "RS" in right:
                    replace_text_in_cell(row.cells[1], {"RS": data['buyer_account']})
                elif "Bank" in right:
                    replace_text_in_cell(row.cells[1], {"Bank": data['buyer_bank']})
                elif "Kor" in right:
                    replace_text_in_cell(row.cells[1], {"Kor": data['buyer_correspondent']})
                elif "Bis" in right:
                    replace_text_in_cell(row.cells[1], {"Bis": data['buyer_bik']})


    invoice_title_para = None
    for paragraph in doc.paragraphs:
        if "СЧЕТ №" in paragraph.text:
            new_text = f"СЧЕТ № {data['invoice_number']} от {data['invoice_date']}"
            replace_text_in_paragraph(paragraph, {paragraph.text: new_text})
            invoice_title_para = paragraph
            break


    services_table = find_table_by_headers(doc, ["№", "Товар / Услуги", "Цена", "Кол-во", "Ед. изм.", "НДС", "Сумма"])
    if services_table:
        tbl = services_table._tbl
        rows_to_remove = []
        for i, row in enumerate(services_table.rows):
            if i == 0:
                continue
            rows_to_remove.append(row._tr)
        for tr in rows_to_remove:
            tbl.remove(tr)

        total_sum = 0
        for idx, service in enumerate(data['services'], start=1):
            new_row = services_table.add_row()
            new_row.cells[0].text = str(idx)
            new_row.cells[1].text = service['name']
            new_row.cells[2].text = f"{service['price']:.2f}"
            new_row.cells[3].text = str(service['quantity'])
            new_row.cells[4].text = service.get('unit', 'шт')
            new_row.cells[5].text = service.get('vat', 'без НДС')
            sum_val = service['price'] * service['quantity']
            new_row.cells[6].text = f"{sum_val:.2f}"
            total_sum += sum_val

        total_row = services_table.add_row()
        if len(total_row.cells) >= 7:
            total_row.cells[0].merge(total_row.cells[5])
            total_words = rubles_to_text(total_sum)
            total_row.cells[0].text = f"Итого (без НДС): {total_words}"
            total_row.cells[6].text = f"{total_sum:.2f}"


        agreement_row = services_table.add_row()
        if len(agreement_row.cells) >= 7:
            agreement_row.cells[0].merge(agreement_row.cells[-1])
            cell = agreement_row.cells[0]
            for paragraph in cell.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            # Добавляем два параграфа (без пустого):
            p1 = cell.add_paragraph()
            p1.text = f"Оплата по договору № {data['contract_number']} от {data['contract_date']} за Услуги клининга (без НДС)"
            p2 = cell.add_paragraph()
            p2.text = "Назначение платежа"
            for run in p2.runs:
                run.italic = True

    set_document_font(doc, "Times New Roman")


    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Индивидуальный предприниматель"):
            set_paragraph_font(paragraph, "Times New Roman", 12, bold=True)
            break


    if invoice_title_para:
        set_paragraph_font(invoice_title_para, "Times New Roman", 12, bold=True)


    if services_table:
        for i, row in enumerate(services_table.rows):
            if i == 0:
                continue
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

                if i == len(services_table.rows) - 2:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    print(f"Документ сохранён: {output_path}")
